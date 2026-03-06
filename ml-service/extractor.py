"""
Resume text extraction from PDF and DOCX files.
"""
import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {str(e)}")

    if not text.strip():
        raise ValueError(
            "No text found in PDF. The file may be scanned/image-based. "
            "Please upload a text-based PDF."
        )
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {str(e)}")

    if not text.strip():
        raise ValueError("No text found in DOCX file.")
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to correct extractor based on file extension."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif fname.endswith(".docx") or fname.endswith(".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Use PDF or DOCX.")